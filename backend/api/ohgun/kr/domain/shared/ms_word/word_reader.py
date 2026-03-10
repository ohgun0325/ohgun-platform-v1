"""python-docx 기반 Word(.docx) 문서 읽기."""

from __future__ import annotations

from pathlib import Path
from typing import List, Dict, Any, Union

try:
    from docx import Document
except ImportError:
    raise ImportError("python-docx가 필요합니다. pip install python-docx")

from app.domain.shared.ms_word.base import WordExtractionStrategy


class DocxReader(WordExtractionStrategy):
    """python-docx를 사용한 .docx 파일 읽기."""

    def extract_text(self, source: Union[str, Path]) -> str:
        """문서 전체 텍스트를 한 문자열로 반환합니다.

        Args:
            source: .docx 파일 경로

        Returns:
            추출된 텍스트
        """
        doc = Document(str(source))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    def extract_paragraphs(self, source: Union[str, Path]) -> List[str]:
        """단락(paragraph) 리스트를 반환합니다.

        Args:
            source: .docx 파일 경로

        Returns:
            단락 문자열 리스트
        """
        doc = Document(str(source))
        return [p.text for p in doc.paragraphs if p.text.strip()]

    def extract_tables(self, source: Union[str, Path]) -> List[Dict[str, Any]]:
        """문서 내 표를 추출합니다.

        Args:
            source: .docx 파일 경로

        Returns:
            표 리스트. 각 항목은 {"rows": [[cell, ...], ...], "row_count", "col_count"} 형태
        """
        doc = Document(str(source))
        result = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text.strip() for cell in row.cells])
            result.append({
                "rows": rows,
                "row_count": len(rows),
                "col_count": len(rows[0]) if rows else 0,
            })
        return result

    def extract_all(self, source: Union[str, Path]) -> Dict[str, Any]:
        """텍스트, 단락, 표를 한 번에 추출합니다.

        Args:
            source: .docx 파일 경로

        Returns:
            {"text", "paragraphs", "tables"} 딕셔너리
        """
        return {
            "text": self.extract_text(source),
            "paragraphs": self.extract_paragraphs(source),
            "tables": self.extract_tables(source),
        }


def read_docx_text(path: Union[str, Path]) -> str:
    """Word 문서에서 텍스트만 추출하는 헬퍼 함수."""
    return DocxReader().extract_text(path)


def read_docx_tables(path: Union[str, Path]) -> List[Dict[str, Any]]:
    """Word 문서에서 표만 추출하는 헬퍼 함수."""
    return DocxReader().extract_tables(path)
